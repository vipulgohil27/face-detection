# resume_reader.py
import docx


# def load_resume_docx(path):
#     doc = docx.Document(path)
#     return "\n".join([para.text for para in doc.paragraphs])
#
#
# def save_resume_docx(text, path):
#     doc = docx.Document()
#     for line in text.split("\n"):
#         doc.add_paragraph(line)
#     doc.save(path)
#

from docx import Document

def load_resume_docx(path):
    doc = Document(path)
    return "\n".join([para.text for para in doc.paragraphs])

def save_resume_docx(text, path):
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    doc.save(path)

